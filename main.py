import importlib.util
import shutil
import subprocess
import sys

_REQUIRED_PIP = [
    "fastapi", "uvicorn", "openai", "pypdf",
    "ebooklib", "python-docx", "beautifulsoup4",
    "fpdf2", "python-multipart",
]
_IMPORT_NAMES = {
    "python-docx": "docx", "beautifulsoup4": "bs4", "fpdf2": "fpdf",
}

_missing = []
for pkg in _REQUIRED_PIP:
    import_name = _IMPORT_NAMES.get(pkg, pkg)
    if importlib.util.find_spec(import_name) is None:
        _missing.append(pkg)

if _missing:
    print(f"[启动] 检测到缺失依赖: {', '.join(_missing)}，正在使用清华镜像安装…", flush=True)
    uv_path = shutil.which("uv")
    cmd, display = ([uv_path, "pip", "install", "-i", "https://pypi.tuna.tsinghua.edu.cn/simple"] + _missing,
                    f"{uv_path} pip install") if uv_path else \
                   ([sys.executable, "-m", "pip", "install", "-i", "https://pypi.tuna.tsinghua.edu.cn/simple"] + _missing,
                    f"{sys.executable} -m pip install")
    print(f"[启动] 执行: {display} …", flush=True)
    result = subprocess.run(cmd, capture_output=False, text=True)
    if result.returncode != 0:
        print(f"[启动] 安装失败（exit code {result.returncode}）", flush=True)
        print("[启动] 请手动执行:", " ".join(cmd), flush=True)
        sys.exit(1)
    print("[启动] 依赖安装完成！", flush=True)

import asyncio
import base64
import io
import json
import os
import re
import time
import uuid
from pathlib import Path
from typing import Optional

from fastapi import FastAPI, Request, UploadFile, File, Form
from fastapi.responses import JSONResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from openai import AsyncOpenAI
from pydantic import BaseModel

# ---------------------------------------------------------------------------
# FastAPI application
# ---------------------------------------------------------------------------

app = FastAPI()

# ---------------------------------------------------------------------------
# Progress persistence
# ---------------------------------------------------------------------------

PROGRESS_DIR = Path(__file__).parent / "progress"
PROGRESS_DIR.mkdir(exist_ok=True)

# ---------------------------------------------------------------------------
# Cancel events — task_id → asyncio.Event
# ---------------------------------------------------------------------------

cancel_events: dict[str, asyncio.Event] = {}

# ---------------------------------------------------------------------------
# API key persistence (obfuscated, local file)
# ---------------------------------------------------------------------------

_OBFUSCATION_KEY = b"TranslatorTool2024!"
_KEY_FILE = Path(__file__).parent / ".translator_key"


def _obfuscate(key: str) -> str:
    data = key.encode("utf-8")
    xored = bytes(
        data[i] ^ _OBFUSCATION_KEY[i % len(_OBFUSCATION_KEY)]
        for i in range(len(data))
    )
    return base64.b64encode(xored).decode("ascii")


def _deobfuscate(obfuscated: str) -> str:
    try:
        xored = base64.b64decode(obfuscated.encode("ascii"))
        data = bytes(
            xored[i] ^ _OBFUSCATION_KEY[i % len(_OBFUSCATION_KEY)]
            for i in range(len(xored))
        )
        return data.decode("utf-8")
    except Exception:
        return ""

# ---------------------------------------------------------------------------
# In-memory task store for polling-based endpoint (legacy compat)
# ---------------------------------------------------------------------------

tasks: dict[str, dict] = {}

# ---------------------------------------------------------------------------
# Document converter
# ---------------------------------------------------------------------------


class DocConverter:
    """Convert uploaded documents to plain UTF-8 text."""

    SUPPORTED = {".txt", ".pdf", ".epub", ".docx"}

    @staticmethod
    def convert(content: bytes, ext: str) -> str:
        ext = ext.lower().lstrip(".")
        if ext == "txt":
            return content.decode("utf-8", errors="replace")
        if ext == "pdf":
            return DocConverter._pdf(content)
        if ext == "epub":
            return DocConverter._epub(content)
        if ext == "docx":
            return DocConverter._docx(content)
        raise ValueError(f"不支持的文件格式: .{ext}")

    @staticmethod
    def _pdf(content: bytes) -> str:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
        if not pages:
            raise ValueError("PDF 文件无法提取到任何文本内容")
        return "\n\n".join(pages)

    @staticmethod
    def _epub(content: bytes) -> str:
        import ebooklib
        from bs4 import BeautifulSoup
        from ebooklib import epub

        book = epub.read_epub(io.BytesIO(content))
        texts: list[str] = []
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), "html.parser")
                for tag in soup.find_all(["p", "h1", "h2", "h3", "h4", "h5", "h6"]):
                    t = tag.get_text(strip=True)
                    if t:
                        texts.append(t)
        if not texts:
            raise ValueError("EPUB 文件无法提取到任何文本内容")
        return "\n\n".join(texts)

    @staticmethod
    def _docx(content: bytes) -> str:
        import docx
        doc = docx.Document(io.BytesIO(content))
        paras: list[str] = []
        for p in doc.paragraphs:
            t = p.text.strip()
            if t:
                paras.append(t)
        if not paras:
            raise ValueError("DOCX 文件无法提取到任何文本内容")
        return "\n\n".join(paras)


# ---------------------------------------------------------------------------
# Text preprocessing & splitting
# ---------------------------------------------------------------------------

_SENTENCE_BOUNDARY = re.compile(r"(?<=[.!?。！？；;])\s+")


def preprocess_text(text: str) -> str:
    """Normalise line endings and collapse excessive blank lines."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{4,}", "\n\n\n", text)
    return text.strip()


def split_text_with_meta(text: str, max_chars: int = 1500) -> list[dict]:
    """Split text into chunks with paragraph-boundary metadata.

    Returns a list of ``{"text": str, "is_paragraph_start": bool}``.
    """
    paragraphs = re.split(r"\n\n+", text.strip())
    result: list[dict] = []
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        sub_chunks = _split_chunk(para, max_chars)
        for i, sc in enumerate(sub_chunks):
            result.append({"text": sc, "is_paragraph_start": i == 0})
    return result


def split_text(text: str, max_chars: int = 1500) -> list[str]:
    """Split text into chunks (legacy, no metadata)."""
    return [c["text"] for c in split_text_with_meta(text, max_chars)]


def _split_chunk(text: str, max_chars: int) -> list[str]:
    if len(text) <= max_chars:
        return [text]
    sentences = _SENTENCE_BOUNDARY.split(text)
    result: list[str] = []
    buf = ""
    for sent in sentences:
        if not sent.strip():
            continue
        if len(buf) + len(sent) + 1 <= max_chars:
            buf = (buf + " " + sent.strip()).strip() if buf else sent.strip()
        else:
            if buf:
                result.append(buf)
            if len(sent) > max_chars:
                result.extend(_hard_split(sent, max_chars))
                buf = ""
            else:
                buf = sent.strip()
    if buf:
        result.append(buf)
    return result


def _hard_split(text: str, max_chars: int) -> list[str]:
    parts: list[str] = []
    pos = 0
    while pos < len(text):
        end = min(pos + max_chars, len(text))
        if end < len(text):
            cut = end
            while cut > pos and text[cut] not in " \t\n.,!?;:":
                cut -= 1
            if cut == pos:
                cut = end
            else:
                cut += 1
            parts.append(text[pos:cut].strip())
            pos = cut
        else:
            parts.append(text[pos:end].strip())
            break
    return [p for p in parts if p]


# ---------------------------------------------------------------------------
# Async chunk translator (uses Semaphore for concurrency control)
# ---------------------------------------------------------------------------


def _build_prompt(chunk: str, source_lang: str, target_lang: str) -> str:
    if source_lang == "auto":
        return (
            f"请将以下内容翻译成{target_lang}，"
            f"只输出翻译结果，不要有任何额外解释：\n\n{chunk}"
        )
    return (
        f"请将以下{source_lang}翻译成{target_lang}，"
        f"只输出翻译结果，不要有任何额外解释：\n\n{chunk}"
    )


async def _translate_chunk_async(
    idx: int,
    chunk: str,
    api_key: str,
    base_url: str,
    model: str,
    source_lang: str,
    target_lang: str,
    enable_search: bool,
    semaphore: asyncio.Semaphore,
    cancel_event: asyncio.Event,
    max_attempts: int = 3,
) -> tuple:
    """Translate one chunk via AsyncOpenAI.

    Returns ``(idx, text, elapsed, attempts, error_msg)``.
    ``text`` is ``None`` when cancelled or all retries exhausted.
    """
    if cancel_event.is_set():
        return (idx, None, 0.0, 0, None)

    async with semaphore:
        if cancel_event.is_set():
            return (idx, None, 0.0, 0, None)

        client = AsyncOpenAI(api_key=api_key, base_url=base_url)
        prompt = _build_prompt(chunk, source_lang, target_lang)

        last_error: str | None = None
        start = time.time()

        for attempt in range(max_attempts):
            try:
                kwargs = {
                    "model": model,
                    "messages": [{"role": "user", "content": prompt}],
                    "stream": False,
                }
                if enable_search:
                    kwargs["extra_body"] = {"enable_search": True}

                resp = await client.chat.completions.create(**kwargs)
                elapsed = time.time() - start
                content = resp.choices[0].message.content
                text = content.strip() if content else chunk
                return (idx, text, elapsed, attempt + 1, None)

            except Exception as exc:
                last_error = str(exc)
                if attempt < max_attempts - 1:
                    await asyncio.sleep(1)

        elapsed = time.time() - start
        return (idx, None, elapsed, max_attempts, last_error)


# ---------------------------------------------------------------------------
# Progress file helpers
# ---------------------------------------------------------------------------


def _save_progress(
    task_id: str,
    translated: dict[int, str],
    total: int,
    params: dict,
):
    """Write current progress to ``progress/<task_id>.json``."""
    data = {
        "task_id": task_id,
        "total_chunks": total,
        "translated": {str(k): v for k, v in translated.items()},
        # Store params *except* API key
        "params": {
            k: v for k, v in params.items() if k != "api_key"
        },
    }
    path = PROGRESS_DIR / f"{task_id}.json"
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _load_progress(task_id: str) -> dict | None:
    path = PROGRESS_DIR / f"{task_id}.json"
    if not path.exists():
        return None
    with open(path, encoding="utf-8") as f:
        return json.load(f)


# ---------------------------------------------------------------------------
# SSE helpers
# ---------------------------------------------------------------------------


def _sse_event(event_type: str, data: dict) -> str:
    return f"event: {event_type}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


async def _sse_generate(params: dict):
    """Async generator: doc convert → preprocess → chunk → parallel translate
    (asyncio + Semaphore).  Pushes SSE events:
      - task_id, status, log, progress, complete / cancelled / error.
    """
    task_id = uuid.uuid4().hex[:12]
    cancel_event = asyncio.Event()
    cancel_events[task_id] = cancel_event

    try:
        # Emit task_id immediately so the frontend can wire up cancel.
        yield _sse_event("task_id", {"task_id": task_id})

        api_key = params["api_key"]
        model = params["model"]
        base_url = params.get("base_url", "https://api.deepseek.com/v1")
        max_workers = params["max_concurrency"]
        source_lang = params["source_language"]
        target_lang = params["target_language"]
        enable_search = params.get("enable_search", False)
        raw_text = params.get("text")
        file_b64 = params.get("file_content")
        file_name = params.get("file_name")
        resume_data: dict | None = params.get("resume_data")

        loop = asyncio.get_running_loop()

        # ---- Resolve source text ----
        if file_b64:
            yield _sse_event("status", {"message": "正在解析文档…"})
            file_bytes = base64.b64decode(file_b64)
            ext = Path(file_name or "unknown.txt").suffix.lower()
            raw_text = await loop.run_in_executor(
                None, DocConverter.convert, file_bytes, ext
            )

        if not raw_text or not raw_text.strip():
            yield _sse_event("error", {"error": "文本内容为空"})
            return

        # ---- Preprocess ----
        raw_text = preprocess_text(raw_text)

        # ---- Chunking with paragraph-boundary metadata ----
        chunks_meta = split_text_with_meta(raw_text)
        chunks = [c["text"] for c in chunks_meta]
        total = len(chunks)
        if total == 0:
            yield _sse_event("error", {"error": "文本内容为空"})
            return

        yield _sse_event(
            "status", {"message": f"文本已分为 {total} 块，开始翻译…"}
        )

        # ---- Determine which chunks are already translated (resume) ----
        translated: dict[int, str] = {}
        skip_set: set[int] = set()

        if resume_data:
            for k, v in resume_data.get("translated", {}).items():
                idx = int(k)
                translated[idx] = v
                skip_set.add(idx)
            source_lang = resume_data.get("params", {}).get(
                "source_lang", source_lang
            )
            target_lang = resume_data.get("params", {}).get(
                "target_lang", target_lang
            )
            enable_search = resume_data.get("params", {}).get(
                "enable_search", enable_search
            )
            yield _sse_event(
                "status",
                {
                    "message": (
                        f"恢复进度: {len(skip_set)}/{total} 块已翻译，"
                        f"继续剩余 {total - len(skip_set)} 块…"
                    )
                },
            )
            # Yield saved chunks immediately so the UI fills in quickly
            for i in range(total):
                if i in translated:
                    yield _sse_event("progress", {
                        "chunk_id": i,
                        "total": total,
                        "translated_text": translated[i],
                        "new_paragraph": chunks_meta[i]["is_paragraph_start"],
                    })

        # ---- Parallel translation with Semaphore ----
        semaphore = asyncio.Semaphore(max_workers)
        success_count = len(skip_set)
        error_count = 0
        start_time = time.time()

        async def worker(idx: int, chunk: str):
            if idx in skip_set:
                return
            result = await _translate_chunk_async(
                idx, chunk, api_key, base_url, model,
                source_lang, target_lang,
                enable_search, semaphore, cancel_event,
            )
            return result

        work_tasks = {
            asyncio.create_task(worker(i, c)): i
            for i, c in enumerate(chunks)
        }
        pending = set(work_tasks.keys())

        while pending and not cancel_event.is_set():
            done_set, pending = await asyncio.wait(
                pending, return_when=asyncio.FIRST_COMPLETED
            )
            for t in done_set:
                idx, text, elapsed, attempts, err = t.result()
                if text is not None:
                    translated[idx] = text
                    success_count += 1

                    yield _sse_event("progress", {
                        "chunk_id": idx,
                        "total": total,
                        "translated_text": text,
                        "new_paragraph": chunks_meta[idx]["is_paragraph_start"],
                    })
                    yield _sse_event("log", {
                        "message": (
                            f"第 {idx} 块完成  "
                            f"耗时 {elapsed:.1f}s  尝试 {attempts} 次"
                        ),
                        "level": "info",
                    })

                    # Auto-save every 10 chunks
                    completed = len(translated)
                    if completed % 10 == 0:
                        _save_progress(
                            task_id, translated, total, params
                        )
                        yield _sse_event("log", {
                            "message": (
                                f"自动保存完成  "
                                f"{completed}/{total} 块"
                            ),
                            "level": "info",
                        })
                else:
                    error_count += 1
                    yield _sse_event("log", {
                        "message": (
                            f"第 {idx} 块失败  "
                            f"耗时 {elapsed:.1f}s  错误: {err}"
                        ),
                        "level": "error",
                    })

        # Cancel remaining on user abort
        if cancel_event.is_set():
            for t in pending:
                t.cancel()
            # Final save
            _save_progress(task_id, translated, total, params)
            yield _sse_event("cancelled", {
                "message": "翻译已终止，进度已保存至本地",
                "task_id": task_id,
                "completed": len(translated),
                "total": total,
            })
            return

        # ---- All done ----
        # One last save so the final state is persisted
        _save_progress(task_id, translated, total, params)

        elapsed = round(time.time() - start_time, 1)
        yield _sse_event("complete", {
            "total_chunks": total,
            "success_count": success_count,
            "error_count": error_count,
            "total_time": elapsed,
            "success_rate": round(success_count / total * 100, 1),
        })

    except asyncio.CancelledError:
        pass
    except Exception as exc:
        yield _sse_event("error", {"error": str(exc)})
    finally:
        cancel_events.pop(task_id, None)


# ---------------------------------------------------------------------------
# API key persistence endpoints
# ---------------------------------------------------------------------------


@app.get("/api/key")
async def get_api_key():
    """Return the saved (obfuscated) API key and base URL, if any."""
    if _KEY_FILE.exists():
        content = _KEY_FILE.read_text("utf-8").strip()
        try:
            data = json.loads(content)
            key = _deobfuscate(data.get("key", ""))
            base_url = data.get("base_url", "https://api.deepseek.com/v1")
            return {"key": key, "base_url": base_url}
        except (json.JSONDecodeError, KeyError):
            pass
    return {"key": "", "base_url": "https://api.deepseek.com/v1"}


@app.post("/api/key")
async def save_api_key(request: Request):
    """Save API key and base URL to local file (key obfuscated)."""
    body = await request.json()
    key = body.get("key", "")
    base_url = body.get("base_url", "https://api.deepseek.com/v1")
    data = {
        "key": _obfuscate(key) if key else "",
        "base_url": base_url,
    }
    if key:
        _KEY_FILE.write_text(json.dumps(data, ensure_ascii=False), encoding="utf-8")
    return {"status": "saved"}


# ---------------------------------------------------------------------------
# Model list query (OpenAI-compatible /v1/models)
# ---------------------------------------------------------------------------


@app.get("/api/models")
async def get_models(base_url: str = "", api_key: str = ""):
    """Fetch available models from the provider's ``/models`` endpoint."""
    if not base_url or not api_key:
        return {"models": []}
    try:
        from openai import OpenAI as SyncOpenAI
        client = SyncOpenAI(api_key=api_key, base_url=base_url)
        resp = client.models.list()
        models = sorted([m.id for m in resp])
        return {"models": models}
    except Exception:
        return {"models": []}


# ---------------------------------------------------------------------------
# API endpoints — translation
# ---------------------------------------------------------------------------


class TranslateRequest(BaseModel):
    api_key: str
    model: str = "deepseek-chat"
    text: str
    max_concurrency: int = 5
    base_url: str = "https://api.deepseek.com/v1"


async def _stream_translate_from_deepseek(client, model: str, text: str):
    messages = [
        {
            "role": "user",
            "content": f"请将以下英文翻译成中文，只输出翻译结果：{text}",
        }
    ]
    stream = await client.chat.completions.create(
        model=model, messages=messages, stream=True,
    )
    async for chunk in stream:
        delta = chunk.choices[0].delta
        if delta.content:
            yield delta.content


@app.post("/translate")
async def translate(req: TranslateRequest):
    """Simple text translation (kept from v1)."""
    client = AsyncOpenAI(
        api_key=req.api_key, base_url=req.base_url
    )
    chunks: list[str] = []
    async for chunk in _stream_translate_from_deepseek(
        client, req.model, req.text
    ):
        chunks.append(chunk)
    return {"result": "".join(chunks)}


@app.post("/translate_file")
async def translate_file(
    file: UploadFile = File(...),
    api_key: str = Form(...),
    model: str = Form("deepseek-chat"),
    max_concurrency: int = Form(5),
    source_language: str = Form("auto"),
    target_language: str = Form("中文"),
):
    """Legacy polling-based file translation."""
    ext = Path(file.filename or "unknown").suffix.lower()
    if ext not in DocConverter.SUPPORTED:
        return JSONResponse(
            status_code=400,
            content={
                "error": (
                    f"不支持的文件格式: {ext}，"
                    f"支持: {', '.join(sorted(DocConverter.SUPPORTED))}"
                )
            },
        )

    content = await file.read()

    from concurrent.futures import ThreadPoolExecutor

    async def _run_legacy_task(tid: str):
        task = tasks[tid]
        try:
            task["status"] = "converting"
            loop = asyncio.get_running_loop()
            text = await loop.run_in_executor(
                None, DocConverter.convert, task["file_content"], task["file_ext"]
            )
            chunks = split_text(text)
            total = len(chunks)
            task["progress"]["total"] = total
            task["total_chunks"] = total
            if total == 0:
                task["status"] = "done"
                task["result"] = ""
                task["end_time"] = time.time()
                return
            task["status"] = "translating"
            executor = ThreadPoolExecutor(max_workers=task["max_concurrency"])
            from openai import OpenAI as SyncOpenAI
            sync_client = SyncOpenAI(
                api_key=task["api_key"], base_url="https://api.deepseek.com/v1"
            )
            def _sync_translate(chunk: str) -> str:
                prompt = _build_prompt(
                    chunk, task["source_lang"], task["target_lang"]
                )
                for att in range(3):
                    try:
                        r = sync_client.chat.completions.create(
                            model=task["model"],
                            messages=[{"role": "user", "content": prompt}],
                        )
                        t = r.choices[0].message.content
                        return t.strip() if t else chunk
                    except Exception:
                        if att < 2:
                            time.sleep(1)
                return f"[翻译失败]"
            from concurrent.futures import as_completed
            futs = {executor.submit(_sync_translate, c): i for i, c in enumerate(chunks)}
            translated = {}
            sc = ec = 0
            for f in as_completed(futs):
                idx = futs[f]
                try:
                    translated[idx] = f.result()
                    sc += 1
                except Exception:
                    translated[idx] = ""
                    ec += 1
                task["progress"]["done"] = sc + ec
            executor.shutdown(wait=False)
            ordered = [translated[i] for i in sorted(translated)]
            task["result"] = "\n\n".join(ordered)
            task["success_count"] = sc
            task["error_count"] = ec
            task["status"] = "done"
        except Exception as exc:
            task["status"] = "error"
            task["error"] = str(exc)
        finally:
            task["end_time"] = time.time()

    task_id = uuid.uuid4().hex[:12]
    tasks[task_id] = {
        "status": "pending",
        "progress": {"total": 0, "done": 0},
        "file_content": content,
        "file_ext": ext,
        "api_key": api_key,
        "model": model,
        "max_concurrency": max_concurrency,
        "source_lang": source_language,
        "target_lang": target_language,
        "result": "",
        "start_time": time.time(),
        "end_time": 0.0,
        "total_chunks": 0,
        "success_count": 0,
        "error_count": 0,
        "error": None,
    }

    asyncio.create_task(_run_legacy_task(task_id))
    return {"task_id": task_id}


@app.get("/progress/{task_id}")
async def get_progress(task_id: str):
    """Poll endpoint for legacy file translation progress."""
    task = tasks.get(task_id)
    if task is None:
        return {"status": "not_found", "error": "任务不存在"}

    resp: dict = {
        "status": task["status"],
        "progress": task["progress"],
        "elapsed": round(time.time() - task["start_time"], 1),
    }
    if task["status"] in ("done", "error"):
        resp["result"] = task["result"]
        resp["total_time"] = round(task["end_time"] - task["start_time"], 1)
        resp["total_chunks"] = task["total_chunks"]
        resp["success_count"] = task["success_count"]
        resp["error_count"] = task["error_count"]
        resp["success_rate"] = (
            round(task["success_count"] / task["total_chunks"] * 100, 1)
            if task["total_chunks"] > 0
            else 0
        )
    if task["status"] == "error":
        resp["error"] = task["error"]
    return resp


# ---------------------------------------------------------------------------
# SSE streaming endpoint
# ---------------------------------------------------------------------------


class StreamTranslateRequest(BaseModel):
    api_key: str
    model: str = "deepseek-chat"
    max_concurrency: int = 5
    source_language: str = "auto"
    target_language: str = "中文"
    base_url: str = "https://api.deepseek.com/v1"
    text: Optional[str] = None
    file_content: Optional[str] = None
    file_name: Optional[str] = None
    enable_search: bool = False


@app.post("/stream_translate")
async def stream_translate(req: StreamTranslateRequest):
    """Translate text or document, streaming per-chunk SSE results."""
    if not req.text and not req.file_content:
        return JSONResponse(
            status_code=400,
            content={"error": "请提供 text 或 file_content 参数"},
        )
    return StreamingResponse(
        _sse_generate(req.model_dump()),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ---------------------------------------------------------------------------
# Resume & cancel endpoints
# ---------------------------------------------------------------------------


@app.post("/resume_translate/{task_id}")
async def resume_translate(task_id: str, request: Request):
    """Resume a previously cancelled/interrupted translation."""
    body = await request.json()
    api_key = body.get("api_key", "")
    if not api_key:
        return JSONResponse(
            status_code=400,
            content={"error": "请提供 API Key"},
        )

    progress = _load_progress(task_id)
    if progress is None:
        return JSONResponse(
            status_code=404,
            content={"error": "进度文件不存在"},
        )

    params = progress.get("params", {})
    params["api_key"] = api_key
    params["resume_data"] = progress

    return StreamingResponse(
        _sse_generate(params),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.get("/check_progress/{task_id}")
async def check_progress(task_id: str):
    """Check whether a progress file exists (for resume UI)."""
    prog = _load_progress(task_id)
    if prog is None:
        return {"found": False}
    translated = prog.get("translated", {})
    return {
        "found": True,
        "total_chunks": prog.get("total_chunks", 0),
        "completed": len(translated),
    }


@app.post("/cancel/{task_id}")
async def cancel_translation(task_id: str):
    """Signal cancellation of an in-flight SSE translation."""
    ev = cancel_events.get(task_id)
    if ev is None:
        return {"status": "not_found"}
    ev.set()
    return {"status": "cancelling"}


# ---------------------------------------------------------------------------
# Export endpoints
# ---------------------------------------------------------------------------

_CHINESE_FONT_CANDIDATES: list[tuple[str, bool]] = [
    ("/System/Library/Fonts/PingFang.ttc", True),
    ("/System/Library/Fonts/STSong.ttf", False),
    ("/Library/Fonts/Arial Unicode.ttf", False),
    ("C:\\Windows\\Fonts\\simsun.ttc", True),
    ("C:\\Windows\\Fonts\\msyh.ttc", True),
    ("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc", True),
    ("/usr/share/fonts/truetype/wqy/wqy-microhei.ttc", True),
    ("/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc", True),
    ("/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc", True),
]


def _find_chinese_font() -> tuple[str, bool] | None:
    for path, is_ttc in _CHINESE_FONT_CANDIDATES:
        if Path(path).exists():
            return path, is_ttc
    return None


@app.post("/export/txt")
async def export_txt(request: Request):
    text = (await request.body()).decode("utf-8")
    ts = int(time.time())
    filename = f"translated_book_{ts}.txt"
    return Response(
        content=text,
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/export/docx")
async def export_docx(request: Request):
    text = (await request.body()).decode("utf-8")

    from docx import Document
    from docx.shared import Pt
    from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.makeelement(qn("w:rFonts"), {})
    rFonts.set(qn("w:ascii"), "宋体")
    rFonts.set(qn("w:hAnsi"), "宋体")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rFonts.set(qn("w:cs"), "宋体")
    rPr.insert(0, rFonts)

    for line in text.split("\n"):
        stripped = line.strip()
        if stripped:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    ts = int(time.time())
    filename = f"translated_book_{ts}.docx"

    return Response(
        content=buf.read(),
        media_type=(
            "application/vnd.openxmlformats-"
            "officedocument.wordprocessingml.document"
        ),
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.post("/export/pdf")
async def export_pdf(request: Request):
    text = (await request.body()).decode("utf-8")

    from fpdf import FPDF

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=20)

    font_found = _find_chinese_font()
    if font_found:
        font_path, is_ttc = font_found
        kwargs = {"fname": font_path, "uni": True}
        if is_ttc:
            kwargs["index"] = 0
        pdf.add_font("CJK", "", **kwargs)
        pdf.set_font("CJK", "", 12)
    else:
        pdf.set_font("Helvetica", "", 12)

    for line in text.split("\n"):
        stripped = line.strip()
        if stripped:
            pdf.multi_cell(0, 6, stripped)
        else:
            pdf.ln(3)

    ts = int(time.time())
    filename = f"translated_book_{ts}.pdf"

    return Response(
        content=pdf.output(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Serve static frontend (last so API routes take precedence)
# ---------------------------------------------------------------------------

app.mount("/", StaticFiles(directory="static", html=True), name="static")

if __name__ == "__main__":
    import uvicorn
    print(f"🌐 打开浏览器访问: http://localhost:8000")
    uvicorn.run("main:app", host="0.0.0.0", port=8000, reload=True)
